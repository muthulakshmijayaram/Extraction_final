import streamlit as st
from google.generativeai import GenerativeModel
import google.generativeai as genai
from docx import Document
import io
from dotenv import load_dotenv
import os

# Load environment variables from .env file
load_dotenv()

# Configure your Google API Key
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")  # Get API key from environment variable
genai.configure(api_key=GOOGLE_API_KEY)

# Initialize the model
model = GenerativeModel('gemini-2.0-flash')

# Function to summarize PDF page by page
def summarize_pdf_by_page(file):
    try:
        pdf_data = file.read()
        pdf_blob = {
            "mime_type": "application/pdf",
            "data": pdf_data
        }
        # Ask Gemini to return content page by page
        message_parts = [
            "Extract   the content from each page of the PDF. For each page, start with 'Page X' (where X is the page number), then the content. Separate each page with two newlines.Read line by line ans extract if menaing unmatched then understand the meaning and extract the content.Read all the content and extract the content accuartely",
            pdf_blob
        ]
        response = model.generate_content(message_parts)
        return response.text
    except Exception as e:
        return f"Error: {str(e)}"

# Streamlit UI
st.set_page_config(page_title="PDF Content Extractor", layout="centered")
st.title("📄 PDF Extractor ")

uploaded_files = st.file_uploader("Upload your PDFs", type=["pdf"], accept_multiple_files=True)

if "extracted_results" not in st.session_state:
    st.session_state["extracted_results"] = {}

if uploaded_files:
    for uploaded_file in uploaded_files:
        st.markdown(f"#### File: {uploaded_file.name}")

        # Use file name as cache key
        cache_key = uploaded_file.name

        if cache_key in st.session_state["extracted_results"]:
            result = st.session_state["extracted_results"][cache_key]
        else:
            with st.spinner(f"Analyzing {uploaded_file.name}..."):
                result = summarize_pdf_by_page(uploaded_file)
            st.session_state["extracted_results"][cache_key] = result

        if result.startswith("Error"):
            st.error(result)
        else:
            st.success("Extraction successful!")
            st.text_area(f"Extracted Content ({uploaded_file.name})", result, height=300)

            # --- DOCX Download Button ---
            doc = Document()
            # Split by double newlines for pages
            pages = [p.strip() for p in result.split('\n\n') if p.strip()]
            for page in pages:
                # Add page number as heading if present, else as normal paragraph
                lines = page.split('\n', 1)
                if lines and lines[0].strip().lower().startswith("page"):
                    doc.add_heading(lines[0].strip(), level=2)
                    if len(lines) > 1:
                        doc.add_paragraph(lines[1].strip())
                else:
                    doc.add_paragraph(page)
            docx_bytes = io.BytesIO()
            doc.save(docx_bytes)
            docx_bytes.seek(0)
            st.download_button(
                label=f"⬇️ Download as DOCX ({uploaded_file.name})",
                data=docx_bytes,
                file_name=f"{uploaded_file.name.rsplit('.',1)[0]}_extracted.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

